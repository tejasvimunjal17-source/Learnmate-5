"""
backend/resume_generator.py
------------------------------
Turns a backend.resume_store.ResumeProfile into an ATS-friendly resume
file (PDF or DOCX), returned as bytes for st.download_button().

Education section format
-------------------------
ResumeProfile.education is a single free-text field (no dedicated
dataclass), populated by frontend/resume_builder.py's dynamic education
cards as blocks like:

    Bachelor of Commerce (Honours) | Pursuing
    Dronacharya Government College
    Gurugram University
    2025 – 2029          Gurugram, Haryana, India

_parse_education_blocks() below segments that free text back into
structured (title, years, institution, location, board, grade) blocks -
anchored on the line that contains a 4-digit year (the "years + location"
line) - and both renderers lay each block out as:

    Degree | Status .......................... 2025 - 2029   (right-aligned)
    Institution ............................... Location       (right-aligned, italic)
    University / Board                                          (bold + italic)

This parsing is self-contained to this module; ResumeProfile and
save_resume() are untouched.

ATS notes
----------
Right-aligned dates require either a table or a right tab stop - a
single-row, borderless 2-column table (PDF) and a right tab stop (DOCX,
no table at all) are used for exactly that purpose. Everything else keeps
the original single-column, bullet-based, no-table layout.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_TAB_ALIGNMENT

from backend.resume_store import ResumeProfile
from backend.logger_setup import get_logger

logger = get_logger(__name__)


class ResumeGenerationError(RuntimeError):
    """Raised when a resume file (PDF or DOCX) fails to generate."""


def _validate_profile(profile: ResumeProfile) -> None:
    if not isinstance(profile, ResumeProfile):
        raise ResumeGenerationError(f"Expected a ResumeProfile instance, got {type(profile).__name__}.")
    if not profile.first_name or not profile.first_name.strip():
        raise ResumeGenerationError("Cannot generate a resume without first_name.")
    if not profile.last_name or not profile.last_name.strip():
        raise ResumeGenerationError("Cannot generate a resume without last_name.")


def _contact_line(profile: ResumeProfile) -> str:
    return profile.email.strip() if profile.email else ""


# ------------------------------------------------------------------
# Education parsing (free text -> structured blocks)
# ------------------------------------------------------------------
# Anchor line: contains a 4-digit year, optionally followed by 2+ spaces
# and a location (e.g. "2025 – 2029          Gurugram, Haryana, India",
# or "2025          Gurugram, Haryana, India" for a single passing year).
_YEAR_LOCATION_RE = re.compile(r"^(?P<years>.*?\d{4}.*?)(?:\s{2,}(?P<location>\S.*))?$")


@dataclass
class _EducationBlock:
    title: str = ""
    years: str = ""
    institution: str = ""
    location: str = ""
    board: str = ""
    grade: str = ""


def _parse_education_blocks(education_text: str) -> list[_EducationBlock]:
    """Segment ResumeProfile.education free text into structured blocks."""
    if not education_text or not education_text.strip():
        return []

    lines = [l for l in education_text.splitlines() if l.strip()]
    raw_blocks: list[list[str]] = []
    buffer: list[str] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()
        buffer.append(line)

        if stripped.lower().startswith("grade:"):
            raw_blocks.append(buffer)
            buffer = []
            i += 1
            continue

        if len(buffer) >= 2 and _YEAR_LOCATION_RE.match(stripped):
            if i + 1 < n and lines[i + 1].strip().lower().startswith("grade:"):
                buffer.append(lines[i + 1])
                i += 2
            else:
                i += 1
            raw_blocks.append(buffer)
            buffer = []
            continue

        i += 1

    if buffer:
        raw_blocks.append(buffer)  # leftover lines that didn't match the pattern

    blocks: list[_EducationBlock] = []
    for raw in raw_blocks:
        block = _EducationBlock()
        remaining = list(raw)

        if remaining and remaining[-1].strip().lower().startswith("grade:"):
            block.grade = remaining[-1].split(":", 1)[-1].strip()
            remaining = remaining[:-1]

        if remaining:
            m = _YEAR_LOCATION_RE.match(remaining[-1].strip())
            if m:
                block.years = (m.group("years") or "").strip()
                block.location = (m.group("location") or "").strip()
                remaining = remaining[:-1]

        if remaining:
            block.title = remaining[0].strip()
        if len(remaining) >= 2:
            block.institution = remaining[1].strip()
        if len(remaining) >= 3:
            block.board = remaining[2].strip()

        blocks.append(block)

    return blocks


# ------------------------------------------------------------------
# PDF (reportlab)
# ------------------------------------------------------------------
def _add_education_pdf(story: list, blocks: list[_EducationBlock], heading_style, content_width: float) -> None:
    title_style = ParagraphStyle("EduTitle", fontName="Helvetica-Bold", fontSize=10, leading=14, alignment=TA_LEFT)
    year_style = ParagraphStyle("EduYear", fontName="Helvetica", fontSize=10, leading=14, alignment=TA_RIGHT)
    inst_style = ParagraphStyle("EduInst", fontName="Helvetica", fontSize=10, leading=14, alignment=TA_LEFT)
    loc_style = ParagraphStyle("EduLoc", fontName="Helvetica-Oblique", fontSize=10, leading=14, alignment=TA_RIGHT)
    board_style = ParagraphStyle("EduBoard", fontName="Helvetica-BoldOblique", fontSize=10, leading=14, alignment=TA_LEFT)
    grade_style = ParagraphStyle("EduGrade", fontName="Helvetica", fontSize=9, leading=12,
                                  alignment=TA_LEFT, textColor="#444444")

    story.append(Paragraph("EDUCATION", heading_style))
    left_w = content_width * 0.62
    right_w = content_width - left_w

    for idx, b in enumerate(blocks):
        rows = []
        if b.title or b.years:
            rows.append([Paragraph(b.title, title_style), Paragraph(b.years, year_style)])
        if b.institution or b.location:
            rows.append([Paragraph(b.institution, inst_style), Paragraph(b.location, loc_style)])

        if rows:
            table = Table(rows, colWidths=[left_w, right_w])
            table.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(table)

        if b.board:
            story.append(Paragraph(b.board, board_style))
        if b.grade:
            story.append(Paragraph(f"Grade: {b.grade}", grade_style))

        if idx < len(blocks) - 1:
            story.append(Spacer(1, 6))


def build_resume_pdf(profile: ResumeProfile) -> bytes:
    """Generate an ATS-friendly, single-column PDF resume.

    Args:
        profile: The ResumeProfile to render.

    Returns:
        The generated PDF file's raw bytes.

    Raises:
        ResumeGenerationError: if `profile` is invalid or PDF rendering fails.
    """
    _validate_profile(profile)

    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=LETTER,
            topMargin=0.6 * inch, bottomMargin=0.6 * inch,
            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
            title=f"{profile.full_name} - Resume",
        )
        styles = getSampleStyleSheet()
        name_style = ParagraphStyle(
            "NameStyle", parent=styles["Title"], fontName="Helvetica-Bold",
            fontSize=18, alignment=TA_LEFT, spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "ContactStyle", parent=styles["Normal"], fontSize=9.5,
            textColor="#333333", spaceAfter=10,
        )
        heading_style = ParagraphStyle(
            "SectionHeading", parent=styles["Heading2"], fontName="Helvetica-Bold",
            fontSize=11.5, spaceBefore=10, spaceAfter=4, textColor="#1a1a1a",
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=3,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=body_style, leftIndent=12, bulletIndent=0,
        )

        story = [
            Paragraph(profile.full_name or "Your Name", name_style),
        ]
        contact = _contact_line(profile)
        if contact:
            story.append(Paragraph(contact, contact_style))
        else:
            story.append(Spacer(1, 6))

        def add_heading(text: str) -> None:
            story.append(Paragraph(text.upper(), heading_style))

        if profile.education:
            education_blocks = _parse_education_blocks(profile.education)
            if education_blocks:
                _add_education_pdf(story, education_blocks, heading_style, doc.width)
            else:
                add_heading("Education")
                for line in profile.education.splitlines():
                    if line.strip():
                        story.append(Paragraph(f"- {line.strip()}", bullet_style))

        if profile.skills:
            add_heading("Skills")
            story.append(Paragraph(", ".join(profile.skills), body_style))

        if profile.internships:
            add_heading("Internship Experience")
            for i in profile.internships:
                header = f"<b>{i.role}</b>"
                if i.company:
                    header += f" - {i.company}"
                if i.duration:
                    header += f" ({i.duration})"
                story.append(Paragraph(header, body_style))
                if i.description:
                    story.append(Paragraph(f"- {i.description}", bullet_style))

        if profile.projects:
            add_heading("Projects")
            for p in profile.projects:
                title = p.title + (f" ({p.tech_stack})" if p.tech_stack else "")
                story.append(Paragraph(f"<b>{title}</b>", body_style))
                if p.description:
                    story.append(Paragraph(f"- {p.description}", bullet_style))

        if profile.certificates:
            add_heading("Certifications")
            for c in profile.certificates:
                label = c.name
                if c.issuer:
                    label += f" - {c.issuer}"
                if c.year:
                    label += f" ({c.year})"
                story.append(Paragraph(f"- {label}", bullet_style))

        if profile.achievements:
            add_heading("Achievements")
            for line in profile.achievements.splitlines():
                if line.strip():
                    story.append(Paragraph(f"- {line.strip()}", bullet_style))

        if profile.hobbies:
            add_heading("Hobbies")
            story.append(Paragraph(", ".join(profile.hobbies), body_style))

        story.append(Spacer(1, 4))
        doc.build(story)
        return buffer.getvalue()

    except ResumeGenerationError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to generate PDF resume for %s", getattr(profile, "email", "unknown"))
        raise ResumeGenerationError(f"Could not generate PDF resume: {exc}") from exc


# ------------------------------------------------------------------
# DOCX (python-docx)
# ------------------------------------------------------------------
def _add_education_docx(doc: Document, blocks: list[_EducationBlock], add_heading) -> None:
    section = doc.sections[0]
    usable_width_in = Emu(
        section.page_width.emu - section.left_margin.emu - section.right_margin.emu
    ).inches

    add_heading("Education")

    for idx, b in enumerate(blocks):
        if b.title or b.years:
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width_in), WD_TAB_ALIGNMENT.RIGHT)
            title_run = p.add_run(b.title)
            title_run.bold = True
            if b.years:
                p.add_run(f"\t{b.years}")

        if b.institution or b.location:
            p2 = doc.add_paragraph()
            p2.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width_in), WD_TAB_ALIGNMENT.RIGHT)
            p2.add_run(b.institution)
            if b.location:
                loc_run = p2.add_run(f"\t{b.location}")
                loc_run.italic = True

        if b.board:
            p3 = doc.add_paragraph()
            board_run = p3.add_run(b.board)
            board_run.bold = True
            board_run.italic = True

        if b.grade:
            grade_p = doc.add_paragraph(f"Grade: {b.grade}")
            if grade_p.runs:
                grade_p.runs[0].font.size = Pt(9)

        if idx < len(blocks) - 1:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)


def build_resume_docx(profile: ResumeProfile) -> bytes:
    """Generate an ATS-friendly, single-column DOCX resume.

    Args:
        profile: The ResumeProfile to render.

    Returns:
        The generated DOCX file's raw bytes.

    Raises:
        ResumeGenerationError: if `profile` is invalid or DOCX rendering fails.
    """
    _validate_profile(profile)

    try:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        for section in doc.sections:
            section.top_margin = section.bottom_margin = Pt(36)
            section.left_margin = section.right_margin = Pt(50)

        name_p = doc.add_paragraph()
        name_run = name_p.add_run(profile.full_name or "Your Name")
        name_run.bold = True
        name_run.font.size = Pt(18)

        contact = _contact_line(profile)
        if contact:
            contact_p = doc.add_paragraph(contact)
            if contact_p.runs:
                contact_p.runs[0].font.size = Pt(9.5)

        def add_heading(text: str) -> None:
            h = doc.add_paragraph()
            run = h.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(12)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(2)

        def add_bullet(text: str) -> None:
            doc.add_paragraph(f"- {text}")

        if profile.education:
            education_blocks = _parse_education_blocks(profile.education)
            if education_blocks:
                _add_education_docx(doc, education_blocks, add_heading)
            else:
                add_heading("Education")
                for line in profile.education.splitlines():
                    if line.strip():
                        add_bullet(line.strip())

        if profile.skills:
            add_heading("Skills")
            doc.add_paragraph(", ".join(profile.skills))

        if profile.internships:
            add_heading("Internship Experience")
            for i in profile.internships:
                p = doc.add_paragraph()
                header = i.role
                if i.company:
                    header += f" - {i.company}"
                if i.duration:
                    header += f" ({i.duration})"
                p.add_run(header).bold = True
                if i.description:
                    add_bullet(i.description)

        if profile.projects:
            add_heading("Projects")
            for proj in profile.projects:
                p = doc.add_paragraph()
                title = proj.title + (f" ({proj.tech_stack})" if proj.tech_stack else "")
                p.add_run(title).bold = True
                if proj.description:
                    add_bullet(proj.description)

        if profile.certificates:
            add_heading("Certifications")
            for c in profile.certificates:
                label = c.name
                if c.issuer:
                    label += f" - {c.issuer}"
                if c.year:
                    label += f" ({c.year})"
                add_bullet(label)

        if profile.achievements:
            add_heading("Achievements")
            for line in profile.achievements.splitlines():
                if line.strip():
                    add_bullet(line.strip())

        if profile.hobbies:
            add_heading("Hobbies")
            doc.add_paragraph(", ".join(profile.hobbies))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ResumeGenerationError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to generate DOCX resume for %s", getattr(profile, "email", "unknown"))
        raise ResumeGenerationError(f"Could not generate DOCX resume: {exc}") from exc
