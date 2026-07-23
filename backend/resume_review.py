"""
backend/resume_review.py
---------------------------
ATS Resume Review backend.

Extracts text from an uploaded PDF or DOCX resume, runs a local (no paid
API) analysis pipeline, and returns an ATS score plus missing sections,
strengths, weaknesses, and improvement suggestions. Review history is
saved via the existing backend/sheets_client.py (Sheets-first, local-CSV
fallback - same pattern already used by auth.py and roadmap_engine.py),
to a sheet named "Resume Reviews".

Note: this project's config.py does not define a sheet-name field for
reviews, so the sheet name is a module-level constant here rather than a
new config.py field, per the "do not modify any other file" instruction.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any

from backend.logger_setup import get_logger
from backend.sheets_client import append_row, read_rows

logger = get_logger(__name__)

RESUME_REVIEWS_SHEET = "Resume Reviews"
REVIEWS_SHEET_HEADER: list[str] = [
    "email", "date", "score", "missing_sections", "strengths", "weaknesses", "suggestions",
]


# ------------------------------------------------------------------
# Text extraction
# ------------------------------------------------------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from an uploaded PDF's bytes using pypdf."""
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.error("Failed to open PDF for text extraction: %s", exc)
        return ""

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - a single bad page shouldn't kill the read
            logger.warning("Failed to extract text from a PDF page; skipping it.")
    return "\n".join(pages).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from an uploaded DOCX's bytes using python-docx."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.error("Failed to open DOCX for text extraction: %s", exc)
        return ""

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


# ------------------------------------------------------------------
# Reference data
# ------------------------------------------------------------------
SECTION_PATTERNS: dict[str, list[str]] = {
    "Contact Information": [r"email", r"phone", r"@", r"linkedin"],
    "Education": [r"education", r"university", r"college", r"degree", r"b\.?tech", r"bachelor", r"master"],
    "Skills": [r"skills", r"technical skills", r"proficienc"],
    "Experience": [r"experience", r"internship", r"employment", r"work history"],
    "Projects": [r"projects?"],
    "Certifications": [r"certificat"],
}

KEYWORD_BANK: list[str] = [
    "python", "sql", "java", "javascript", "react", "node", "aws", "azure",
    "gcp", "cloud", "docker", "kubernetes", "git", "agile", "scrum",
    "machine learning", "data analysis", "communication", "leadership",
    "teamwork", "problem solving", "project management", "api",
    "testing", "ci/cd", "linux", "excel", "tableau", "power bi",
]

WEAK_PHRASES: list[str] = [
    "responsible for", "worked on", "helped with", "duties included",
    "in charge of", "was tasked with",
]

STRONG_VERBS: list[str] = [
    "built", "led", "designed", "implemented", "launched", "optimized",
    "improved", "reduced", "increased", "automated", "delivered",
    "architected", "developed", "created", "managed", "achieved",
]


@dataclass
class ResumeReviewResult:
    ats_score: int
    missing_sections: list[str] = field(default_factory=list)
    strengths: list[str] = field(default_factory=list)
    weaknesses: list[str] = field(default_factory=list)
    suggestions: list[str] = field(default_factory=list)


# ------------------------------------------------------------------
# Section / keyword checks
# ------------------------------------------------------------------
def check_resume_sections(text: str) -> tuple[list[str], list[str]]:
    """Return (present_sections, missing_sections) against SECTION_PATTERNS."""
    lower = text.lower()
    present, missing = [], []
    for section, patterns in SECTION_PATTERNS.items():
        if any(re.search(p, lower) for p in patterns):
            present.append(section)
        else:
            missing.append(section)
    return present, missing


def detect_missing_keywords(text: str) -> tuple[list[str], list[str]]:
    """Return (found_keywords, missing_keywords) against KEYWORD_BANK."""
    lower = text.lower()
    found = [kw for kw in KEYWORD_BANK if kw in lower]
    missing = [kw for kw in KEYWORD_BANK if kw not in lower]
    return found, missing


def _check_formatting(text: str) -> list[str]:
    """Local, dependency-free formatting checks (no external service required)."""
    issues = []
    if "\t" in text:
        issues.append("Tab characters detected - tables/tab layouts can confuse ATS parsers.")

    bullet_lines = sum(1 for l in text.splitlines() if l.strip().startswith(("-", "•", "*")))
    if bullet_lines == 0:
        issues.append("No bullet points detected - use bullets for experience/projects.")

    double_spaces = len(re.findall(r"\S  +\S", text))
    if double_spaces > 3:
        issues.append(f"Found {double_spaces} instances of extra/inconsistent spacing.")

    return issues


def _check_length(text: str) -> list[str]:
    """Local resume-length checks."""
    word_count = len(text.split())
    issues = []
    if word_count < 150:
        issues.append("Resume content looks very short - may read as incomplete to ATS/recruiters.")
    elif word_count > 1200:
        issues.append("Resume content looks very long - consider trimming to 1-2 pages worth of content.")
    return issues


def _check_grammar(text: str) -> list[str]:
    """Lightweight, dependency-free grammar checks (no language_tool_python/server needed)."""
    issues = []
    if re.search(r"(?<!\w)i(?!\w)", text):  # lowercase standalone "i"
        issues.append("Found a lowercase standalone 'i' - should be capitalized 'I'.")
    repeated = re.findall(r"\b(\w+)\s+\1\b", text, flags=re.IGNORECASE)
    if repeated:
        issues.append(f"Possible repeated word detected (e.g. '{repeated[0]}').")
    return issues


def _strengths_and_weaknesses(
    text: str, present_sections: list[str], formatting_issues: list[str],
    length_issues: list[str], grammar_issues: list[str], found_keywords: list[str],
) -> tuple[list[str], list[str]]:
    lower = text.lower()
    strengths, weaknesses = [], []

    weak_hits = sum(1 for phrase in WEAK_PHRASES if phrase in lower)
    strong_hits = sum(1 for verb in STRONG_VERBS if verb in lower)

    if "Experience" in present_sections or "Projects" in present_sections:
        if strong_hits > weak_hits:
            strengths.append("Uses strong action verbs in experience/project descriptions.")
        elif weak_hits > 0:
            weaknesses.append("Relies on passive phrases (e.g. 'responsible for') instead of action verbs.")

    if "Skills" in present_sections:
        strengths.append("Includes a dedicated Skills section.")
    if "Contact Information" in present_sections:
        strengths.append("Contact information is present and easy to find.")
    if len(found_keywords) >= 6:
        strengths.append("Good keyword coverage across common technical/professional terms.")
    elif len(found_keywords) < 3:
        weaknesses.append("Very few recognizable industry keywords found.")

    weaknesses.extend(formatting_issues)
    weaknesses.extend(length_issues)
    weaknesses.extend(grammar_issues)

    if not strengths:
        strengths.append("No standout strengths detected yet - see suggestions below.")
    return strengths, weaknesses


# ------------------------------------------------------------------
# Scoring
# ------------------------------------------------------------------
def calculate_ats_score(
    text: str,
    present_sections: list[str],
    missing_sections: list[str],
    found_keywords: list[str],
) -> int:
    """Weighted 0-100 ATS score from section coverage, keyword density,
    formatting red flags, grammar issues, and resume length."""
    total_sections = len(present_sections) + len(missing_sections)
    section_score = (len(present_sections) / total_sections) * 40 if total_sections else 0

    keyword_score = min(len(found_keywords) / 12, 1.0) * 30

    formatting_issues = _check_formatting(text)
    length_issues = _check_length(text)
    grammar_issues = _check_grammar(text)
    word_count = len(text.split())
    length_score = 20 if 200 <= word_count <= 900 else (12 if 100 <= word_count <= 1200 else 5)

    penalty = min(len(formatting_issues) + len(grammar_issues) + len(length_issues), 4) * 2.5
    quality_score = max(10 - penalty, 0)

    score = section_score + keyword_score + length_score + quality_score
    return max(0, min(100, round(score)))


# ------------------------------------------------------------------
# Suggestions
# ------------------------------------------------------------------
def generate_improvement_tips(
    missing_sections: list[str],
    missing_keywords: list[str],
    weaknesses: list[str],
) -> list[str]:
    """Plain-language improvement suggestions derived from the analysis above."""
    tips: list[str] = []

    if missing_sections:
        tips.append(f"Add these missing sections: {', '.join(missing_sections)}.")
    if missing_keywords:
        sample = ", ".join(missing_keywords[:8])
        tips.append(f"Consider adding relevant keywords if applicable to your field: {sample}.")
    for w in weaknesses:
        if w not in tips:
            tips.append(w if w.endswith((".", "!", "?")) else f"{w}.")

    if not tips:
        tips.append("Resume looks solid overall - only minor polish needed.")
    return tips


# ------------------------------------------------------------------
# Orchestration + persistence
# ------------------------------------------------------------------
def review_resume(text: str, email: str = "") -> ResumeReviewResult:
    """Run the full local review pipeline over extracted resume text."""
    if not text or not text.strip():
        return ResumeReviewResult(
            ats_score=0,
            missing_sections=list(SECTION_PATTERNS.keys()),
            strengths=[],
            weaknesses=["Couldn't extract any text from this file."],
            suggestions=["Please upload a text-based (not scanned/image) PDF or DOCX resume."],
        )

    present_sections, missing_sections = check_resume_sections(text)
    found_keywords, missing_keywords = detect_missing_keywords(text)
    formatting_issues = _check_formatting(text)
    length_issues = _check_length(text)
    grammar_issues = _check_grammar(text)

    strengths, weaknesses = _strengths_and_weaknesses(
        text, present_sections, formatting_issues, length_issues, grammar_issues, found_keywords,
    )
    score = calculate_ats_score(text, present_sections, missing_sections, found_keywords)
    suggestions = generate_improvement_tips(missing_sections, missing_keywords, weaknesses)

    result = ResumeReviewResult(
        ats_score=score,
        missing_sections=missing_sections,
        strengths=strengths,
        weaknesses=weaknesses,
        suggestions=suggestions,
    )

    if email:
        try:
            save_review_history(email, result)
        except Exception:  # noqa: BLE001 - review must succeed even if history save fails
            logger.exception("Failed to save resume review history for %s", email)

    return result


def save_review_history(email: str, result: ResumeReviewResult) -> None:
    """Persist a review summary to the 'Resume Reviews' sheet via the existing sheets_client."""
    row: dict[str, Any] = {
        "email": email,
        "date": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        "score": result.ats_score,
        "missing_sections": ", ".join(result.missing_sections),
        "strengths": " | ".join(result.strengths),
        "weaknesses": " | ".join(result.weaknesses),
        "suggestions": " | ".join(result.suggestions),
    }
    append_row(RESUME_REVIEWS_SHEET, REVIEWS_SHEET_HEADER, row)
    logger.info("Resume review history saved for %s", email)


def get_review_history(email: str) -> list[dict[str, Any]]:
    """Return this user's past resume reviews, most recent last."""
    rows = read_rows(RESUME_REVIEWS_SHEET, REVIEWS_SHEET_HEADER)
    return [r for r in rows if r.get("email") == email]
